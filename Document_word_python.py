# -*- coding: utf-8 -*-
"""
Created on Sat Aug  8 11:27:08 2020

@author: Brenda Rojas Delgado
"""

import docx
import os
from docx.shared import Pt
from docx.shared import Inches
import pandas as pd
from io import BytesIO
import matplotlib.pyplot as plt
import numpy as np
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def Figure(paragraph):
    run = run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = ' SEQ Figure \* ARABIC'
    r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

#Change path to where the data is stored 
path = r'C:\Users\erijo797\Documents\Python Scripts\Word\python_word_document-master\data.xlsx'

## CALCULATIONS
df = pd.read_excel(path)
cases = np.array(df['cases'])
temp = np.array(df['temp'])
## clearing up data, some days no reporting -> remoce instances when cases == 0
temp = np.delete(temp, cases==0)
cases = np.delete(cases, cases==0)
corr = np.corrcoef(temp,cases)[1,0]

##saving scatter plot to memfile
memfile = BytesIO()
plt.scatter(temp,cases)
plt.xlabel("Temperature [\N{DEGREE SIGN}C]")
plt.ylabel("Reported cases")
plt.savefig(memfile)

doc = docx.Document()
style = doc.styles['Normal']
font = style.font
font.size = Pt(12)
main_header=doc.add_heading("Correlation between cases of Covid-19 and temperature", 0)
doc.add_heading("- A case study of Uppsala",0)

doc.add_heading("Introduction",1)
doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. \
Donec bibendum, arcu at hendrerit convallis, nulla turpis \
elementum nisl, nec fermentum mi augue id nisl. Pellentesque \
tempus, quam porta consequat fringilla, magna est maximus \
neque, at venenatis tortor enim eget massa. Sed aliquet \
ultrices tempor. Suspendisse potenti. Donec erat nibh, \
posuere tincidunt metus sit amet, lobortis \
malesuada mi. Sed eget enim molestie, \
ullamcorper mi et, aliquet tellus. \
Vestibulum ac mauris turpis. \
Maecenas porttitor turpis ac suscipit lacinia.")

doc.add_heading("Method", 1)
doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. \
Donec bibendum, arcu at hendrerit convallis, nulla turpis \
elementum nisl, nec fermentum mi augue id nisl. Pellentesque \
tempus, quam porta consequat fringilla, magna est maximus \
neque, at venenatis tortor enim eget massa. Sed aliquet \
ultrices tempor. Suspendisse potenti. Donec erat nibh, \
posuere tincidunt metus sit amet, lobortis \
malesuada mi. Sed eget enim molestie, \
ullamcorper mi et, aliquet tellus. \
Vestibulum ac mauris turpis. \
Maecenas porttitor turpis ac suscipit lacinia.")

doc.add_heading("Results", 1)
doc.add_picture(memfile, width=Inches(5))
paragraph = doc.add_paragraph('Figure 1', style='Caption')
Figure(paragraph)
paragraph.add_run(' Scatter plot of daily average temperatur and reported covid-19 cases ')

doc.add_paragraph("The correllation coefficient between reported number of covid-infections and \
temperature is " + str(round(corr,2)) + " which indicates a low correlation.")

doc.add_heading("Discussion", 1)

doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. \
Donec bibendum, arcu at hendrerit convallis, nulla turpis \
elementum nisl, nec fermentum mi augue id nisl. Pellentesque \
tempus, quam porta consequat fringilla, magna est maximus \
neque, at venenatis tortor enim eget massa. Sed aliquet \
ultrices tempor. Suspendisse potenti. Donec erat nibh, \
posuere tincidunt metus sit amet, lobortis \
malesuada mi. Sed eget enim molestie, \
ullamcorper mi et, aliquet tellus. \
Vestibulum ac mauris turpis. \
Maecenas porttitor turpis ac suscipit lacinia.")

 
doc.save("covidtemp.docx")
os.system("start covidtemp.docx")